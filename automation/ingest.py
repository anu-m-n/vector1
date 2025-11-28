import sys
import os
import uuid
import re
import pandas as pd
from pypdf import PdfReader
from docx import Document
import chromadb
import kuzu

# Fix imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend.embeddings import get_embedding

# --- CONFIGURATION ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")
RAW_FILES_DIR = os.path.join(DATA_DIR, "raw_files")
VECTOR_DB_PATH = os.path.join(DATA_DIR, "chroma_db")
GRAPH_DB_PATH = os.path.join(DATA_DIR, "kuzu_db")

# --- DATABASE SETUP ---
if not os.path.exists(RAW_FILES_DIR): os.makedirs(RAW_FILES_DIR)

print(">>> Connecting to Databases...")
chroma_client = chromadb.PersistentClient(path=VECTOR_DB_PATH)
vector_collection = chroma_client.get_or_create_collection(name="knowledge_base")

db = kuzu.Database(GRAPH_DB_PATH)
conn = kuzu.Connection(db)

# Create Schema
try:
    conn.execute("CREATE NODE TABLE Document(id STRING, title STRING, PRIMARY KEY (id))")
    conn.execute("CREATE NODE TABLE Entity(id STRING, name STRING, PRIMARY KEY (id))")
    conn.execute("CREATE REL TABLE MENTIONS(FROM Document TO Entity)")
except RuntimeError: pass

# --- HELPER FUNCTIONS ---

def parse_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except: return ""

def parse_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    except: return ""

def extract_keywords(text):
    # "Lite" Graph Logic: Find capitalized words (Entities)
    # Ignore common words like 'The', 'This'
    raw_words = re.findall(r'\b[A-Z][a-zA-Z]+\b', text)
    blacklist = {"The", "This", "That", "And", "For", "With", "But"}
    return list(set([w for w in raw_words if w not in blacklist and len(w) > 3]))

def ingest_content(title, text, source_type):
    if len(text) < 10: return

    doc_id = str(uuid.uuid4())
    print(f"   -> Ingesting: {title}...")

    # 1. Vector Store (Mock)
    vector_collection.add(
        documents=[text[:2000]], # Store first 2000 chars
        embeddings=[get_embedding(text)],
        ids=[doc_id],
        metadatas=[{"title": title, "source": source_type}]
    )

    # 2. Graph Store
    safe_title = title.replace("'", "").replace('"', "")
    conn.execute(f"MERGE (d:Document {{id: '{doc_id}', title: '{safe_title}'}})")

    keywords = extract_keywords(text[:5000])
    for word in keywords:
        word_id = word.lower()
        safe_word = word.replace("'", "")
        # Link Document -> Keyword
        conn.execute(f"MERGE (e:Entity {{id: '{word_id}', name: '{safe_word}'}})")
        conn.execute(f"MATCH (d:Document), (e:Entity) WHERE d.id = '{doc_id}' AND e.id = '{word_id}' MERGE (d)-[:MENTIONS]->(e)")

# --- MAIN RUNNER ---

def ingest_all():
    print(f"\n--- Scanning {RAW_FILES_DIR} ---")
    
    files = os.listdir(RAW_FILES_DIR)
    if not files:
        print("No files found! Put PDFs, CSVs, or Text files in 'data/raw_files/'")
        return

    for f in files:
        path = os.path.join(RAW_FILES_DIR, f)
        
        # A. Handle CSVs (Large Databases)
        if f.endswith(".csv"):
            print(f"Processing Database: {f}")
            try:
                df = pd.read_csv(path)
                # Assumes CSV has 'title' and 'overview' columns (Like TMDB)
                # If not, it tries to just read the first text column
                text_col = 'overview' if 'overview' in df.columns else df.columns[1]
                title_col = 'title' if 'title' in df.columns else df.columns[0]
                
                # Load first 100 rows for demo speed
                for _, row in df.head(100).iterrows():
                    ingest_content(str(row[title_col]), str(row[text_col]), "CSV Database")
            except Exception as e:
                print(f"Error reading CSV: {e}")

        # B. Handle Unstructured Files (PDF/Docx)
        elif f.endswith(".pdf"):
            text = parse_pdf(path)
            ingest_content(f, text, "PDF Document")
            
        elif f.endswith(".docx"):
            text = parse_docx(path)
            ingest_content(f, text, "Word Document")
            
        elif f.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as tf:
                ingest_content(f, tf.read(), "Text File")

if __name__ == "__main__":
    ingest_all()
    print("\n>>> INGESTION COMPLETE!")